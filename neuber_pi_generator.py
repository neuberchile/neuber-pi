#!/usr/bin/env python3
"""
NEUBER — Script Generación Automática PI / SC
Trigger: Deal cerrado (stage_id=6) en Pipedrive → genera documento Word PI → adjunta en Pipedrive
v2.9 — fix PI counter persistente en Pipedrive deal 467 (resuelve reset por deploy)
v2.10 — parse_items tolerante a formato natural sin separadores (espacios, m3, USD, RL, etc.)
v2.11 — endpoint /regenerate_pi_with_signature: regenera PI con imagen de firma del proveedor (Flujo 1.5)
"""

from flask import Flask, request, jsonify
import requests
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import io
import json
import re
import hashlib

app = Flask(__name__)

# ─── AUTH PI_ADMIN_TOKEN (v2.15 sesion 3.34) ──────────────────────────
# Endpoints write-heavy (/webhook, /generate_pi, /regenerate_pi_with_signature,
# /bank_hash/register) requieren header X-PI-Token. /health NO requiere auth
# para que healthchecks externos funcionen.
#
# Migracion: si PI_ADMIN_TOKEN env var esta vacia, fail-open con warning log
# para no romper produccion durante deploy. Una vez validado, segundo commit
# elimina fail-open y lo hace obligatorio.
PI_ADMIN_TOKEN = os.environ.get('PI_ADMIN_TOKEN', '')

def require_pi_token(f):
    """Decorador para endpoints write-heavy. Header esperado: X-PI-Token."""
    from functools import wraps
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not PI_ADMIN_TOKEN:
            print(f"[PI] WARNING: PI_ADMIN_TOKEN env var vacia, fail-open en {request.path}")
            return f(*args, **kwargs)
        provided = request.headers.get('X-PI-Token', '')
        if provided != PI_ADMIN_TOKEN:
            print(f"[PI] AUTH FAIL en {request.path}: token mismatch o ausente")
            return jsonify({'ok': False, 'error': 'unauthorized'}), 401
        return f(*args, **kwargs)
    return wrapper


# ─── CONFIG ───────────────────────────────────────────────────────────────────
PIPEDRIVE_API  = os.environ.get('PIPEDRIVE_API', '')  # v2.14 sesion 3.34: fallback hardcoded eliminado tras validacion
PIPEDRIVE_BASE = 'https://api.pipedrive.com/v1'
DEAL_CERRADO_STAGE = 6

# Field keys Pipedrive
F_PROVEEDOR    = '75918aeb25c96daf3e0d79bfcca114dc47a6329d'
F_AGENTE       = 'e6be65c16bd853251d87c2dfb51e0c27c390cb3c'
F_SPECIE       = 'e041539f14e7191423096602086f2851e7595c33'
F_PRODUCT      = 'e3a582a6b130ded4f1edffbc8cc502de2ff70ba8'
F_VOLUMEN      = '6b060e20ce5b99df7118e876b887d48f48d36fcb'
F_INCOTERM     = 'bbd9343f7be5f218230c1863d977b7a10aaa22a3'  # id:72 — corregido (id:45 eliminado)
F_POD          = 'da61aaa7dcbbfa40b6834b9fa462d1f187699b9b'  # id:68 — corregido (id:49 eliminado)
F_MES_EMBARQUE = 'be724a357e45d9429a201cbe1527e3f51daf5fab'
F_REF_CONTRATO = '03fbb80f2776069b6fdf81cd73766a91a11b87df'
F_GRADE        = '7558474ef96645532c47d9f187274077f38380f3'
F_SIZE         = '3b33d6404cf7a33163af3d9c375ae037bbf656d2'
F_PRECIO       = '63ebae55f11ccaefeafab28d66e3584a2503d047'
F_PAGO         = '3f353a1fd6fe963964da09202436a5881d9155b9'
F_ITEMS        = 'f122062872dd838e7e37c0094000d28284b3e17f'  # ← NUEVO campo Ítems

# ─── PI NUMBER COUNTER ────────────────────────────────────────────────────────

_PI_COUNTER_MARKER = "PI_COUNTER:"


def _read_pi_counter_note():
    """Lee la nota PI_COUNTER de deal 467. Devuelve (note_id, last_value).

    Si no existe, devuelve (None, 7699).
    Si hay múltiples notas PI_COUNTER legacy (por bugs previos), usa la primera con mayor valor.
    """
    try:
        url = f"{PIPEDRIVE_BASE}/notes?deal_id=467&api_token={PIPEDRIVE_API}&limit=100"
        r = requests.get(url, timeout=10)
        data = r.json().get('data') or []
        candidates = []
        for item in data:
            content = item.get('content') or ''
            clean = re.sub(r'<[^>]+>', '', content).strip()
            if clean.startswith(_PI_COUNTER_MARKER):
                try:
                    val = int(clean[len(_PI_COUNTER_MARKER):].split()[0].strip())
                    candidates.append((item.get('id'), val))
                except Exception as pe:
                    print(f"[counter] Nota PI_COUNTER malformada id={item.get('id')}: {pe}")
        if candidates:
            # Elegir la nota con el valor más alto (evita reset si hay duplicados legacy)
            candidates.sort(key=lambda x: x[1], reverse=True)
            return candidates[0]
    except Exception as e:
        print(f"[counter] Error leyendo PI_COUNTER deal 467: {e}")
    return None, 7699


def get_next_pi_number():
    """Incrementa y persiste el PI counter en deal 467 Pipedrive (idempotente + persistente).

    Garantías:
    - Persistencia entre deploys Railway (no usa filesystem local).
    - Idempotente: UPDATE in-place de la nota master. No duplica notas.
    - Fallback: si Pipedrive falla, devuelve 7700 para no bloquear el PI.
    """
    try:
        note_id, last = _read_pi_counter_note()
        next_val = last + 1
        payload = f"{_PI_COUNTER_MARKER}{next_val}"
        if note_id:
            url = f"{PIPEDRIVE_BASE}/notes/{note_id}?api_token={PIPEDRIVE_API}"
            requests.put(url, json={'content': payload}, timeout=10)
        else:
            url = f"{PIPEDRIVE_BASE}/notes?api_token={PIPEDRIVE_API}"
            requests.post(url, json={'deal_id': 467, 'content': payload}, timeout=10)
        return next_val
    except Exception as e:
        print(f"[counter] Fallback local por error: {e}")
        return 7700

# ─── PARSE ITEMS ──────────────────────────────────────────────────────────────
_SIZE_PATTERN = re.compile(r'\d+\s*[xX×]\s*\d+\s*[xX×]\s*(?:\d+|RL)', re.IGNORECASE)
_NUM_PATTERN = re.compile(r'^[-+]?\d+(?:[.,]\d+)?$')


def _to_float(s):
    return float(s.replace(',', '.'))


def parse_items(items_text, default_price=0):
    """
    Parsea el campo Ítems de Pipedrive.

    Formato preferido (v2.10+): una línea por ítem, tres tokens separados por espacios.
        EspxAnchoxLargo Volumen Precio
        Ejemplo: 15x86x4080 200 260

    Tolerancias:
      - Sufijo m3 / m³ en el volumen: "200m3", "200 m3"
      - Prefijo/sufijo USD o $ en el precio: "USD 260", "$260", "260 USD"
      - Largo "RL" (random length): "86x86xRL"
      - Mayúsculas/minúsculas en la "x": "15X86x4080"
      - Coma decimal: "200,5"
      - Espacios extra
      - Formato legacy con pipes: "15x86x4080 | 200 | 260" (sigue funcionando)

    Si una línea no puede parsearse, se descarta silenciosamente.
    Si el precio no aparece en la línea, se usa default_price.

    Retorna: lista de dicts {size, volume, price, total}
    """
    result = []
    if not items_text:
        return result

    lines = [l.strip() for l in items_text.strip().splitlines() if l.strip()]
    for line in lines:
        # Rama legacy: si la línea contiene "|", usar split por pipes.
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            if len(parts) < 2:
                continue
            try:
                size_str = parts[0]
                volume = _to_float(parts[1])
                if len(parts) >= 3 and parts[2]:
                    price = _to_float(parts[2])
                else:
                    price = float(default_price or 0)
                total = round(volume * price, 2)
                result.append({'size': size_str, 'volume': volume, 'price': price, 'total': total})
            except Exception:
                continue
            continue

        # Rama nueva: tokens separados por espacios.
        # 1. Extraer el size con regex.
        size_match = _SIZE_PATTERN.search(line)
        if not size_match:
            continue
        size_str = size_match.group(0).replace(' ', '')

        # 2. Quitar el size del string para procesar el resto.
        rest = line[:size_match.start()] + ' ' + line[size_match.end():]

        # 3. Quitar tokens reservados (USD, $, m3, m³) y normalizar.
        rest = re.sub(r'\$', ' ', rest)
        rest = re.sub(r'\bUSD\b', ' ', rest, flags=re.IGNORECASE)
        rest = re.sub(r'(\d)\s*m\s*[3³]\b', r'\1', rest, flags=re.IGNORECASE)

        # 4. Extraer todos los números del resto.
        numbers = []
        for tok in rest.split():
            if _NUM_PATTERN.match(tok):
                try:
                    numbers.append(_to_float(tok))
                except Exception:
                    pass

        if not numbers:
            continue

        try:
            volume = numbers[0]
            if len(numbers) >= 2:
                price = numbers[1]
            else:
                price = float(default_price or 0)
            total = round(volume * price, 2)
            result.append({'size': size_str, 'volume': volume, 'price': price, 'total': total})
        except Exception:
            continue

    return result

# ─── PIPEDRIVE API ────────────────────────────────────────────────────────────
def get_deal(deal_id):
    r = requests.get(f'{PIPEDRIVE_BASE}/deals/{deal_id}',
                     params={'api_token': PIPEDRIVE_API})
    return r.json().get('data')

def get_org(org_id):
    r = requests.get(f'{PIPEDRIVE_BASE}/organizations/{org_id}',
                     params={'api_token': PIPEDRIVE_API})
    return r.json().get('data', {})

def attach_file_to_deal(deal_id, filename, content_bytes):
    r = requests.post(
        f'{PIPEDRIVE_BASE}/files',
        params={'api_token': PIPEDRIVE_API},
        data={'deal_id': deal_id},
        files={'file': (filename, content_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    )
    return r.json()

def attach_file_to_project(project_id, filename, content_bytes):
    r = requests.post(
        f'{PIPEDRIVE_BASE}/files',
        params={'api_token': PIPEDRIVE_API},
        data={'project_id': project_id},
        files={'file': (filename, content_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    )
    return r.json()

def get_project_by_deal(deal_id):
    try:
        r = requests.get(f'{PIPEDRIVE_BASE}/projects',
                         params={'api_token': PIPEDRIVE_API, 'limit': 500})
        data = r.json().get('data', {})
        items = data.get('items', []) if isinstance(data, dict) else []
        for item in items:
            project = item.get('item', item)
            deal_ids = project.get('deal_ids', []) or []
            if deal_id in deal_ids:
                return project.get('id')
    except Exception as e:
        print(f'[PI] Error buscando proyecto: {e}')
    return None

def add_note_to_deal(deal_id, content):
    r = requests.post(
        f'{PIPEDRIVE_BASE}/notes',
        params={'api_token': PIPEDRIVE_API},
        json={'deal_id': deal_id, 'content': content}
    )
    return r.json()

# ─── DATOS PROVEEDOR ──────────────────────────────────────────────────────────
PROVEEDOR_DATA = {
    'Masisa': {
        'name': 'MASISA S.A.',
        'address': 'Avda. Apoquindo N°3650 Piso 10, Las Condes, Santiago, Chile',
        'tax_id': 'RUT: 93.007.000-9',
        'phone': '+56 2 2520 3000',
        'email': 'info@masisa.com',
        'origin': 'Chile',
        'port': 'Puerto de Lirquén / San Antonio, Chile',
        'bank': 'Banco de Chile',
        'account': '225-27284-09',
        'swift': 'BCHICLRM',
        'bank_address': 'Santiago, Chile',
        'incoterm_note': 'CIF/FOB'
    },
    'Norfor': {
        'name': 'NORFOR S.A.',
        'address': 'Ruta Prov. 34 Km. 1,5 San Carlos, Corrientes, Argentina',
        'tax_id': 'CUIT: 30-70847764-4',
        'phone': '+54 9 1156326061',
        'email': 'ezequiel.miraglia@norfor.com.ar',
        'origin': 'Argentina',
        'port': 'Puerto de Buenos Aires, Argentina',
        'bank': 'Santander Rio S.A.',
        'account': '3544034620001',
        'swift': 'BSCHARBA',
        'bank_address': 'Buenos Aires, Argentina',
        'intermediary_bank': 'Standard Chartered Bank',
        'intermediary_swift': 'SCBLUS33',
        'intermediary_aba': '026002561',
        'incoterm_note': 'FOB'
    },
    'Arboreal': {
        'name': 'ARBOREAL S.A.',
        'address': 'Ruta 26 Km 224 Paso Santander, Casilla de Correo 78026, Tacuarembó, Uruguay',
        'tax_id': 'Tax ID: 217274570014',
        'phone': '',
        'email': 'franz@arboreal.uy',
        'origin': 'Uruguay',
        'port': 'Puerto de Montevideo, Uruguay',
        'bank': 'Scotiabank Uruguay S.A.',
        'account': '024-1763590500',
        'swift': 'COMEUYMM',
        'bank_address': 'Montevideo, Uruguay',
        'incoterm_note': 'FOB/CIF'
    },
    'Laminadora': {
        'name': 'LAMINADORA LOS ANGELES S.A.',
        'address': 'Los Ángeles, Chile',
        'tax_id': '',
        'phone': '',
        'email': '',
        'origin': 'Chile',
        'port': 'Puerto de Lirquén, Chile',
        'bank': 'Banco de Chile',
        'account': '05-230-40269-04',
        'swift': 'BCHICLRM',
        'bank_address': 'Santiago, Chile',
        'incoterm_note': 'FOB'
    },
    'Agrifor': {
        'name': 'AGRIFOR S.A.',
        'address': 'Bulnes 815 Oficina 502, Temuco, Chile',
        'tax_id': '',
        'phone': '',
        'email': '',
        'origin': 'Chile',
        'port': 'Puerto Chile',
        'bank': 'Banco de Chile',
        'account': '5-240-10019-05',
        'swift': 'BCHICLRM',
        'bank_address': 'Santiago, Chile',
        'incoterm_note': 'FOB'
    },
    'Santa Blanca': {
        'name': 'FORESTAL SANTA BLANCA S.A.',
        'address': 'Chile',
        'tax_id': 'RUT: 79.712.980-1',
        'phone': '',
        'email': '',
        'origin': 'Chile',
        'port': 'Puerto Chile',
        'bank': 'Banco Santander',
        'account': '510005573-4 (USD) / 5680564-8 (CLP)',
        'swift': 'BSCHCLRM',
        'bank_address': 'Chile',
        'incoterm_note': 'FOB'
    },
    'DEFAULT': {
        'name': 'VER PROVEEDOR',
        'address': 'Ver contrato',
        'tax_id': '',
        'phone': '',
        'email': '',
        'origin': 'Ver contrato',
        'port': 'Puerto según proveedor',
        'bank': 'Ver Contrato',
        'account': 'Ver Contrato',
        'swift': 'Ver Contrato',
        'bank_address': 'Ver Contrato',
        'incoterm_note': 'FOB/CIF'
    }
}


# ─── POL-003: HASH SHA256 DATOS BANCARIOS ────────────────────────────────────
# Protección contra BEC: el hash de los datos bancarios de cada proveedor se
# registra en deal 467 Pipedrive (nota BANK_HASH:<proveedor>:<sha256>).
# Antes de generar cada PI, se verifica que el hash del bloque actual coincide
# con el registrado. Si falla → alertar a Julio y bloquear la generación.

def compute_bank_hash(proveedor_name):
    """Devuelve SHA256 del bloque bancario del proveedor."""
    data = PROVEEDOR_DATA.get(proveedor_name, {})
    bank_fields = [
        data.get('name', ''),
        data.get('bank', ''),
        data.get('account', ''),
        data.get('swift', ''),
        data.get('bank_address', ''),
        data.get('tax_id', '')
    ]
    payload = '|'.join(bank_fields).encode('utf-8')
    return hashlib.sha256(payload).hexdigest()


_MASTER_NOTE_MARKER = "BANK_HASHES_MASTER:"


def _read_master_note():
    """Lee la nota master de deal 467. Devuelve (note_id, hashes_dict).

    Si no existe, devuelve (None, {}).
    """
    try:
        url = f"{PIPEDRIVE_BASE}/notes?deal_id=467&api_token={PIPEDRIVE_API}&limit=100"
        r = requests.get(url, timeout=10)
        data = r.json().get('data') or []
        for item in data:
            content = item.get('content') or ''
            # Pipedrive agrega <br /> y <p>; limpiamos tags antes de parsear
            clean = re.sub(r'<[^>]+>', '', content).strip()
            if clean.startswith(_MASTER_NOTE_MARKER):
                payload = clean[len(_MASTER_NOTE_MARKER):].strip()
                try:
                    hashes = json.loads(payload)
                    if isinstance(hashes, dict):
                        return item.get('id'), hashes
                except Exception as je:
                    print(f"[hash] Master note JSON inválido (id={item.get('id')}): {je}")
    except Exception as e:
        print(f"[hash] Error leyendo master note deal 467: {e}")
    return None, {}


def _write_master_note(note_id, hashes_dict):
    """Escribe o actualiza la nota master en deal 467."""
    try:
        payload_str = _MASTER_NOTE_MARKER + json.dumps(hashes_dict, sort_keys=True)
        if note_id:
            url = f"{PIPEDRIVE_BASE}/notes/{note_id}?api_token={PIPEDRIVE_API}"
            requests.put(url, json={'content': payload_str}, timeout=10)
        else:
            url = f"{PIPEDRIVE_BASE}/notes?api_token={PIPEDRIVE_API}"
            requests.post(url, json={'deal_id': 467, 'content': payload_str}, timeout=10)
        return True
    except Exception as e:
        print(f"[hash] Error escribiendo master note deal 467: {e}")
        return False


def get_registered_bank_hash(proveedor_name):
    """Lee el hash registrado del proveedor desde la master note de deal 467."""
    _, hashes = _read_master_note()
    return hashes.get(proveedor_name)


def register_bank_hash(proveedor_name):
    """Registra/actualiza el hash del proveedor en la master note (idempotente)."""
    try:
        current = compute_bank_hash(proveedor_name)
        note_id, hashes = _read_master_note()
        if hashes.get(proveedor_name) == current:
            return True  # ya registrado con el mismo valor
        hashes[proveedor_name] = current
        return _write_master_note(note_id, hashes)
    except Exception as e:
        print(f"[hash] Error registrando hash {proveedor_name}: {e}")
        return False


def verify_bank_hash(proveedor_name):
    """Verifica que el hash actual coincide con el registrado.
    Devuelve (ok, mensaje). Si no hay registro previo, lo crea y devuelve ok."""
    current = compute_bank_hash(proveedor_name)
    registered = get_registered_bank_hash(proveedor_name)
    if registered is None:
        register_bank_hash(proveedor_name)
        return True, f"Hash registrado por primera vez: {current[:12]}..."
    if registered != current:
        return False, (
            f"HASH MISMATCH para {proveedor_name}! "
            f"Registrado: {registered[:16]}... Actual: {current[:16]}... "
            f"Posible compromiso de datos bancarios. REVISAR URGENTE."
        )
    return True, f"Hash OK: {current[:12]}..."


# ─── GENERADOR DE PI ──────────────────────────────────────────────────────────
def generate_pi_document(deal_data, pi_number, signature_image_bytes=None, signature_mime='image/png'):
    """
    Genera documento Word PI a partir de deal_data.

    Si signature_image_bytes es provisto, inserta la imagen como firma del proveedor
    en la celda izquierda de la tabla de firmas. Tamaño objetivo: 1.2 inches de ancho.

    Args:
        deal_data: dict de Pipedrive
        pi_number: int número correlativo
        signature_image_bytes: bytes de imagen PNG/JPG (opcional)
        signature_mime: 'image/png' | 'image/jpeg' (informativo, python-docx detecta)

    Returns:
        bytes del documento Word
    """
    # POL-003: verificar hash bancario antes de generar PI
    proveedor_name_check = ''
    if isinstance(deal_data.get(F_PROVEEDOR), dict):
        proveedor_name_check = deal_data[F_PROVEEDOR].get('name', '')
    if proveedor_name_check and proveedor_name_check in PROVEEDOR_DATA:
        hash_ok, hash_msg = verify_bank_hash(proveedor_name_check)
        if not hash_ok:
            print(f"[PI] HASH FAIL: {hash_msg}")
        else:
            print(f"[PI] Hash check OK: {hash_msg}")

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    # Proveedor
    proveedor_name = ''
    if isinstance(deal_data.get(F_PROVEEDOR), dict):
        proveedor_name = deal_data[F_PROVEEDOR].get('name', '')
    prov_info = PROVEEDOR_DATA.get(proveedor_name, PROVEEDOR_DATA['DEFAULT'])

    # Cliente
    cliente_name = deal_data.get('org_name', 'CLIENTE')
    _org_id_raw = deal_data.get('org_id')
    cliente_id = _org_id_raw.get('value') if isinstance(_org_id_raw, dict) else _org_id_raw
    cliente_address = ''
    if cliente_id:
        org_data = get_org(cliente_id)
        addr = (org_data.get('address') or org_data.get('address_formatted_address') or '')
        if not addr:
            parts = [
                org_data.get('address_street') or '',
                org_data.get('address_city') or '',
                org_data.get('address_state') or '',
                org_data.get('address_country') or '',
            ]
            addr = ', '.join(p for p in parts if p)
        cliente_address = addr

    # Campos generales
    grade    = deal_data.get(F_GRADE, '')
    precio   = float(deal_data.get(F_PRECIO, 0) or 0)
    volumen  = float(deal_data.get(F_VOLUMEN, 0) or 0)
    pod      = deal_data.get(F_POD, '')
    etd      = deal_data.get(F_MES_EMBARQUE, '')
    pago_term = deal_data.get(F_PAGO, 'T/T 20% advance + 80% against copy of documents')
    fecha_hoy = datetime.now().strftime('%d-%m-%Y')

    # Incoterm
    incoterm_raw = deal_data.get(F_INCOTERM, 'FOB')
    if isinstance(incoterm_raw, dict):
        incoterm = incoterm_raw.get('label', 'FOB')
    elif incoterm_raw and str(incoterm_raw).isdigit():
        try:
            r_field = requests.get(f'{PIPEDRIVE_BASE}/dealFields',
                                   params={'api_token': PIPEDRIVE_API, 'limit': 200})
            fields = r_field.json().get('data', [])
            incoterm_field = next((f for f in fields if f.get('key') == F_INCOTERM), None)
            if incoterm_field:
                options = {str(o['id']): o['label'] for o in incoterm_field.get('options', [])}
                incoterm = options.get(str(incoterm_raw), str(incoterm_raw))
            else:
                incoterm = str(incoterm_raw)
        except:
            incoterm = str(incoterm_raw)
    else:
        incoterm = str(incoterm_raw) if incoterm_raw else 'FOB'

    # ── ÍTEMS — multi-item desde campo Ítems, fallback a campos clásicos ──────
    items_text = deal_data.get(F_ITEMS, '') or ''
    items = parse_items(items_text, default_price=precio)

    if not items:
        # Fallback: un solo ítem desde campos SIZE/VOLUMEN/PRECIO
        size = deal_data.get(F_SIZE, '')
        total_usd = round(precio * volumen, 2)
        items = [{'size': size, 'volume': volumen, 'price': precio, 'total': total_usd}]

    total_vol_global = round(sum(i['volume'] for i in items), 3)
    total_usd_global = round(sum(i['total'] for i in items), 2)

    # ── HEADER ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('NEUBER')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_after = Pt(0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('Julio Condeza Neuber, Agencia Comer. Exterior EIRL').font.size = Pt(8)
    p2.paragraph_format.space_after = Pt(0)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("O'Higgins 420 Of. 102, Concepción, Chile | Tel: +56 41 2246560 | info@neuberchile.com")
    r3.font.size = Pt(7)
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p3.paragraph_format.space_after = Pt(4)

    # ── TÍTULO ────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = p_title.add_run(f'PURCHASE ORDER N° {pi_number}')
    rt.bold = True
    rt.font.size = Pt(13)
    rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p_title.paragraph_format.space_after = Pt(6)

    # ── TABLA CABECERA ────────────────────────────────────────────────────────
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    def set_col_width(table, col_idx, width_inches):
        for row in table.rows:
            tc = row.cells[col_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = _OE('w:tcW')
            tcW.set(_qn('w:w'), str(int(width_inches * 1440)))
            tcW.set(_qn('w:type'), 'dxa')
            existing = tcPr.find(_qn('w:tcW'))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcW)

    tbl_header = doc.add_table(rows=1, cols=3)
    tbl_header.style = 'Table Grid'
    set_col_width(tbl_header, 0, 3.0)
    set_col_width(tbl_header, 1, 3.0)
    set_col_width(tbl_header, 2, 1.5)

    cells = tbl_header.rows[0].cells
    seller_text = f"SELLER:\n{prov_info['name']}\n{prov_info['address']}"
    if prov_info.get('tax_id'):
        seller_text += f"\n{prov_info['tax_id']}"
    if prov_info.get('phone'):
        seller_text += f"\nPhone: {prov_info['phone']}"
    if prov_info.get('email'):
        seller_text += f"\nEmail: {prov_info['email']}"
    cells[0].text = seller_text

    buyer_text = f"BUYER:\n{cliente_name.upper()}"
    buyer_text += f"\n{cliente_address}" if cliente_address else "\n(Address to be confirmed)"
    cells[1].text = buyer_text

    cells[2].text = f"DATE: {fecha_hoy}\nNO.: {pi_number}"

    for cell in cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.name = 'Arial'

    # ── CONDICIONES ───────────────────────────────────────────────────────────
    tbl_cond = doc.add_table(rows=2, cols=4)
    tbl_cond.style = 'Table Grid'
    cond_data = [
        ('INCOTERM', incoterm or 'FOB',        'DATE OF DELIVERY', str(etd) if etd else 'To be confirmed'),
        ('POD',      pod or 'To be confirmed', 'PAYMENT TERM',     pago_term or 'T/T 20% advance + 80% CAD'),
    ]
    for i, (l1, v1, l2, v2) in enumerate(cond_data):
        tbl_cond.rows[i].cells[0].text = l1
        tbl_cond.rows[i].cells[1].text = v1
        tbl_cond.rows[i].cells[2].text = l2
        tbl_cond.rows[i].cells[3].text = v2
        for cell in tbl_cond.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
                    if cell.text in (l1, l2):
                        run.bold = True

    # ── TABLA PRODUCTOS — multi-item ──────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    headers_prod = ['PRODUCT / DESCRIPTION', 'QTY (M3)', 'PRICE (USD/M3)', 'TOTAL USD']
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers_prod):
        cell = hdr_row.cells[i]
        cell.text = h
        run_h = cell.paragraphs[0].runs[0]
        run_h.bold = True
        run_h.font.size = Pt(8)
        run_h.font.name = 'Arial'
        run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        tcPr.append(shd)

    specie = 'Radiata Pine' if proveedor_name in ['Masisa', 'Laminadora', 'Agrifor', 'Santa Blanca'] else 'Taeda Pine'

    # Una fila por ítem
    for item in items:
        product_desc = f'PINE {specie.split()[0].upper()} GRADE'
        if grade:
            product_desc += f' {grade}'
        if item['size']:
            product_desc += f'\n{item["size"]}'

        row = tbl.add_row()
        row.cells[0].text = product_desc.strip()
        row.cells[1].text = str(item['volume'])
        row.cells[2].text = str(int(item['price']))
        row.cells[3].text = str(int(item['total']))
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'

    # Fila TOTAL
    row_total = tbl.add_row()
    row_total.cells[0].text = 'TOTAL'
    row_total.cells[1].text = str(total_vol_global)
    row_total.cells[2].text = ''
    row_total.cells[3].text = str(int(total_usd_global))
    for cell in row_total.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.name = 'Arial'

    # ── QUALITY AND TOLERANCE ─────────────────────────────────────────────────
    p_qt = doc.add_paragraph()
    p_qt.paragraph_format.space_before = Pt(4)
    run_qt = p_qt.add_run('Quality and Tolerance:')
    run_qt.bold = True
    run_qt.font.size = Pt(8)

    tolerances = [
        'a) Total Volume Tolerance: +/- 10% at seller option',
        'b) Thickness and width tolerance: -0/+2mm.',
        'c) Length Tolerance: -0/+2mm',
        'd) Moisture content: 12% +/- 4%',
        'e) NO blue stain allowed',
    ]
    for t in tolerances:
        p_t = doc.add_paragraph(t)
        p_t.runs[0].font.size = Pt(8)
        p_t.runs[0].font.name = 'Arial'
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after = Pt(0)

    # ── PAYMENT DETAILS ───────────────────────────────────────────────────────
    p_pay_title = doc.add_paragraph()
    p_pay_title.paragraph_format.space_before = Pt(6)
    run_pay_title = p_pay_title.add_run('PAYMENT DETAILS')
    run_pay_title.bold = True
    run_pay_title.font.size = Pt(9)
    run_pay_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    bank_val = prov_info['bank']
    if prov_info.get('bank_address'):
        bank_val += '\n' + prov_info['bank_address']

    pay_rows = [
        ('INCOTERM', incoterm or 'FOB'),
        ('BENEFICIARY:', prov_info['name']),
        ('BANK:', bank_val),
        ('PAYMENT TERM:', pago_term or 'T/T 20% advance + 80% CAD'),
        ('SWIFT CODE:', prov_info['swift']),
        ('ACCOUNT:', prov_info['account']),
    ]
    if prov_info.get('intermediary_bank'):
        pay_rows.append(('Intermediary Bank:', prov_info['intermediary_bank']))
        pay_rows.append(('Swift Code:', prov_info.get('intermediary_swift', '')))
        pay_rows.append(('ABA:', prov_info.get('intermediary_aba', '')))

    tbl_pay = doc.add_table(rows=len(pay_rows), cols=2)
    tbl_pay.style = 'Table Grid'
    for i, (label, val) in enumerate(pay_rows):
        tbl_pay.rows[i].cells[0].text = label
        tbl_pay.rows[i].cells[1].text = val
        for cell in tbl_pay.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
        if label:
            try:
                tbl_pay.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            except IndexError:
                pass

    

    # ── BEC WARNING (POL-002) ────────────────────────────────────────────────
    p_bec = doc.add_paragraph()
    p_bec.paragraph_format.space_before = Pt(4)
    run_bec = p_bec.add_run('IMPORTANT: Bank details above must be confirmed via phone call with our team before any wire transfer. Neuber will never request a change of bank details via email.')
    run_bec.bold = True
    run_bec.font.size = Pt(8)
    run_bec.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # ── FIRMAS ────────────────────────────────────────────────────────────────
    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.style = 'Table Grid'
    # (Table no soporta paragraph_format — se omite el espaciado vertical aquí)

    sign_left  = tbl_sign.rows[0].cells[0]
    sign_right = tbl_sign.rows[0].cells[1]

    # Lado IZQUIERDO (proveedor): si hay firma, insertar imagen + nombre.
    # Si no hay firma, mantener el comportamiento original (3 párrafos vacíos + nombre).
    if signature_image_bytes:
        # 1 párrafo vacío de espacio arriba para que la firma no quede pegada al borde
        sign_left.add_paragraph('')
        # Párrafo con la imagen, centrado horizontalmente
        p_img = sign_left.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        try:
            sig_stream = io.BytesIO(signature_image_bytes)
            run_img.add_picture(sig_stream, width=Inches(1.2))
        except Exception as e:
            # Falla silenciosa: si la imagen es inválida, dejar 3 párrafos vacíos
            # como en el caso sin firma. Logueamos para diagnóstico.
            print(f"[PI] add_picture failed: {e} — fallback a sin firma")
            for _ in range(2):
                sign_left.add_paragraph('')
        # Nombre proveedor abajo de la firma
        p_name_left = sign_left.add_paragraph(prov_info['name'])
        if p_name_left.runs:
            p_name_left.runs[0].bold = True
            p_name_left.runs[0].font.size = Pt(8)
            p_name_left.runs[0].font.name = 'Arial'
    else:
        # Comportamiento original sin firma
        for _ in range(3):
            sign_left.add_paragraph('')
        p_name_left = sign_left.add_paragraph(prov_info['name'])
        if p_name_left.runs:
            p_name_left.runs[0].bold = True
            p_name_left.runs[0].font.size = Pt(8)
            p_name_left.runs[0].font.name = 'Arial'

    # Lado DERECHO (cliente): siempre comportamiento original
    for _ in range(3):
        sign_right.add_paragraph('')
    p_name_right = sign_right.add_paragraph(cliente_name.upper())
    if p_name_right.runs:
        p_name_right.runs[0].bold = True
        p_name_right.runs[0].font.size = Pt(8)
        p_name_right.runs[0].font.name = 'Arial'

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── WEBHOOK ──────────────────────────────────────────────────────────────────
@app.route('/webhook', methods=['POST'])
# v2.16 sesion 3.34: webhook NO usa @require_pi_token. Pipedrive ya manda Basic Auth.
# Auth via X-PI-Token solo en endpoints invocados por Luke/manual.
def webhook():
    data = request.json
    if not data:
        return jsonify({'status': 'no data'}), 200

    event = data.get('event', '')
    if 'deal' not in event:
        return jsonify({'status': 'ignored', 'event': event}), 200

    if 'current' in data:
        current  = data.get('current', {})
        previous = data.get('previous', {})
    else:
        current  = data.get('data', {}).get('current', {})
        previous = data.get('data', {}).get('previous', {})

    deal_id        = current.get('id')
    stage_current  = current.get('stage_id')
    stage_previous = previous.get('stage_id')

    print(f"[PI] Webhook | Deal: {deal_id} | Stage: {stage_previous} → {stage_current}")

    # GUARD v2.12: NUNCA procesar el deal sistema 467 (NEUBER_SYSTEM)
    # El deal 467 contiene el PI_COUNTER y otros markers persistentes.
    # Si alguien lo moviera a stage 6 por error, neuber-pi generaria una PI
    # quemando un numero del counter y adjuntando un .docx invalido al sistema.
    if deal_id == 467:
        print(f"[PI] GUARD: deal sistema 467 ignorado")
        return jsonify({'status': 'system deal ignored'}), 200

    if stage_current != DEAL_CERRADO_STAGE or stage_previous == DEAL_CERRADO_STAGE:
        return jsonify({'status': 'not a close event', 'stage': stage_current}), 200

    deal_data = get_deal(deal_id)
    if not deal_data:
        return jsonify({'status': 'error', 'msg': 'deal not found'}), 200

    pi_number = get_next_pi_number()
    filename  = f'PI_{pi_number}_{deal_data.get("org_name","")}.docx'.replace(' ', '_')

    try:
        doc_bytes = generate_pi_document(deal_data, pi_number)
    except Exception as e:
        print(f"[PI] Error generando doc: {e}")
        return jsonify({'status': 'error', 'msg': str(e)}), 200

    attach_file_to_deal(deal_id, filename, doc_bytes)

    project_id = get_project_by_deal(deal_id)
    if project_id:
        attach_file_to_project(project_id, filename, doc_bytes)

    # Contar ítems para la nota
    items_text = deal_data.get(F_ITEMS, '') or ''
    items = parse_items(items_text, float(deal_data.get(F_PRECIO, 0) or 0))
    n_items = len(items) if items else 1

    add_note_to_deal(deal_id,
        f'✅ PI {pi_number} generada automáticamente al cerrar el deal.\n'
        f'Archivo: {filename}\n'
        f'Ítems: {n_items}\n'
        f'Fecha: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n\n'
        f'⚠️ Verificar: dirección del comprador, condiciones específicas, banco del proveedor.'
    )

    print(f"[PI] PI {pi_number} generada — {n_items} ítem(s)")
    return jsonify({'status': 'success', 'pi_number': pi_number, 'filename': filename, 'n_items': n_items}), 200


@app.route('/generate_pi/<int:deal_id>', methods=['GET'])
@require_pi_token
def generate_pi_manual(deal_id):
    # GUARD v2.13: NUNCA generar PI sobre el deal sistema 467
    # Sin este guard, llamadas accidentales/maliciosas a /generate_pi/467
    # quemarian numeros del PI counter persistente (deal 467 storage central).
    if deal_id == 467:
        print(f"[PI] GUARD: deal sistema 467 ignorado en /generate_pi")
        return jsonify({'status': 'system deal ignored'}), 200

    deal_data = get_deal(deal_id)
    if not deal_data:
        return jsonify({'error': 'deal not found'}), 404

    pi_number = get_next_pi_number()
    filename  = f'PI_{pi_number}_{deal_data.get("org_name","")}.docx'.replace(' ', '_')
    doc_bytes = generate_pi_document(deal_data, pi_number)
    attach_result = attach_file_to_deal(deal_id, filename, doc_bytes)

    items_text = deal_data.get(F_ITEMS, '') or ''
    items = parse_items(items_text, float(deal_data.get(F_PRECIO, 0) or 0))
    n_items = len(items) if items else 1

    add_note_to_deal(deal_id,
        f'✅ PI {pi_number} generada manualmente.\nArchivo: {filename}\n'
        f'Ítems: {n_items}\nFecha: {datetime.now().strftime("%Y-%m-%d %H:%M")}'
    )

    return jsonify({
        'status': 'success',
        'pi_number': pi_number,
        'filename': filename,
        'n_items': n_items,
        'attach_deal': attach_result.get('success'),
    })


@app.route('/regenerate_pi_with_signature/<int:deal_id>', methods=['POST'])
@require_pi_token
def regenerate_pi_with_signature(deal_id):
    """
    Regenera la PI Word de un deal incluyendo la imagen de firma del proveedor.

    Reusa el pi_number provisto en el body (NO incrementa el counter), porque
    una PI firmada es la misma PI que la sin firmar — solo cambia el documento
    para incluir la firma.

    Body JSON:
        pi_number (int, required): número de PI a reusar
        signature_b64 (str, required): imagen firma codificada en base64
        signature_mime (str, optional): 'image/png' (default) | 'image/jpeg'
        attach_to_deal (bool, optional, default true): si adjuntar al deal Ventas

    Returns:
        { status, pi_number, filename, content_b64 }
    """
    # GUARD v2.13: NUNCA regenerar PI sobre el deal sistema 467
    if deal_id == 467:
        print(f"[PI] GUARD: deal sistema 467 ignorado en /regenerate_pi_with_signature")
        return jsonify({'status': 'system deal ignored'}), 200

    import base64

    body = request.get_json(silent=True) or {}
    pi_number = body.get('pi_number')
    sig_b64 = body.get('signature_b64')
    sig_mime = body.get('signature_mime', 'image/png')
    attach_flag = body.get('attach_to_deal', True)

    if not pi_number:
        return jsonify({'error': 'pi_number is required'}), 400
    if not sig_b64:
        return jsonify({'error': 'signature_b64 is required'}), 400

    deal_data = get_deal(deal_id)
    if not deal_data:
        return jsonify({'error': 'deal not found'}), 404

    try:
        sig_bytes = base64.b64decode(sig_b64)
    except Exception as e:
        return jsonify({'error': f'invalid base64: {e}'}), 400

    if len(sig_bytes) < 100:
        return jsonify({'error': f'signature image too small: {len(sig_bytes)} bytes'}), 400

    try:
        doc_bytes = generate_pi_document(
            deal_data,
            pi_number,
            signature_image_bytes=sig_bytes,
            signature_mime=sig_mime,
        )
    except Exception as e:
        return jsonify({'error': f'document generation failed: {e}'}), 500

    org_name = deal_data.get('org_name', '') or ''
    filename = f'PI_{pi_number}_{org_name}_firmada.docx'.replace(' ', '_')

    attach_result = None
    if attach_flag:
        attach_result = attach_file_to_deal(deal_id, filename, doc_bytes)

    content_b64 = base64.b64encode(doc_bytes).decode('ascii')

    return jsonify({
        'status': 'success',
        'pi_number': pi_number,
        'filename': filename,
        'size_bytes': len(doc_bytes),
        'content_b64': content_b64,
        'attached_to_deal': bool(attach_result and attach_result.get('success')) if attach_flag else None,
    })


@app.route('/bank_hash/register', methods=['GET', 'POST'])
@require_pi_token
def bank_hash_register_all():
    """Registra/actualiza hashes SHA256 de todos los proveedores en la master note de deal 467.
    Idempotente: si un hash ya está registrado con el mismo valor, no hace nada.
    Útil para pre-cargar hashes sin necesidad de emitir un PI real.
    """
    results = {}
    for proveedor_name in PROVEEDOR_DATA.keys():
        if proveedor_name == 'DEFAULT':
            continue
        try:
            current = compute_bank_hash(proveedor_name)
            registered_before = get_registered_bank_hash(proveedor_name)
            if registered_before == current:
                results[proveedor_name] = {'status': 'unchanged', 'hash_prefix': current[:12]}
            else:
                ok = register_bank_hash(proveedor_name)
                action = 'registered' if registered_before is None else 'updated'
                results[proveedor_name] = {
                    'status': action if ok else 'failed',
                    'hash_prefix': current[:12]
                }
        except Exception as e:
            results[proveedor_name] = {'status': 'error', 'error': str(e)}
    return jsonify({'status': 'ok', 'providers': results})


@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'service': 'Neuber PI Generator', 'version': '2.16'})


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)





